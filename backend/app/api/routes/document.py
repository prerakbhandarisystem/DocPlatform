from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from fastapi.responses import Response
from sqlalchemy.orm import Session
from typing import List
import io
import docx
import PyPDF2
from app.core.database import get_db
from app.models.document import Document
from app.services.google_drive import get_google_drive_service

router = APIRouter()

def extract_text_from_document(file_data: bytes, filetype: str, filename: str) -> str:
    """Extract text content from document based on file type."""
    try:
        if filetype in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'] or filename.endswith('.docx'):
            # Extract text from Word document
            doc = docx.Document(io.BytesIO(file_data))
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            return '\n\n'.join(text_content)
            
        elif filetype == 'application/pdf' or filename.endswith('.pdf'):
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            text_content = []
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(page_text.strip())
            return '\n\n'.join(text_content)
            
        elif filetype.startswith('text/') or filename.endswith(('.txt', '.md')):
            # Plain text file
            return file_data.decode('utf-8', errors='ignore')
            
        else:
            return f"Text extraction not supported for file type: {filetype}"
            
    except Exception as e:
        return f"Error extracting text: {str(e)}"

@router.post("/upload", status_code=201)
async def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    try:
        # Read file data
        file_data = await file.read()
        
        print(f"📁 Uploading {file.filename} ({len(file_data)} bytes)")
        print(f"📄 File type: {file.content_type}")
        
        # Try to upload to Google Drive
        google_drive_id = None
        google_drive_service = get_google_drive_service()
        
        if google_drive_service:
            print("✅ Google Drive service initialized")
            try:
                google_drive_id = google_drive_service.upload_document(
                    file_data, file.filename, file.content_type
                )
                if google_drive_id:
                    print(f"✅ Google Drive upload successful: {google_drive_id}")
                else:
                    print("❌ Google Drive upload failed - no file ID returned")
            except Exception as e:
                print(f"❌ Google Drive upload error: {str(e)}")
        else:
            print("❌ Google Drive service not available")
        
        # Save to database
        document = Document(
            filename=file.filename,
            filetype=file.content_type,
            filesize=len(file_data),
            file_data=file_data,
            google_drive_id=google_drive_id
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        result = {
            "id": document.id,
            "filename": document.filename,
            "google_drive_id": google_drive_id,
            "message": "Document uploaded successfully"
        }
        
        if google_drive_id and google_drive_service:
            result["embed_url"] = google_drive_service.get_embed_url(google_drive_id)
            result["message"] += " with Google Docs integration"
        else:
            result["message"] += " (local storage only - Google Drive not available)"
        
        return result
        
    except Exception as e:
        db.rollback()
        print(f"❌ Upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.get("")
async def get_documents(db: Session = Depends(get_db)):
    """Get all documents."""
    documents = db.query(Document).all()
    return [
        {
            "id": doc.id,
            "filename": doc.filename,
            "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
            "filesize": doc.filesize,
            "filetype": doc.filetype
        }
        for doc in documents
    ]

@router.get("/{document_id}/content")
async def get_document_content(document_id: str, db: Session = Depends(get_db)):
    """Get document content by ID."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return Response(
        content=document.file_data,
        media_type=document.filetype,
        headers={
            "Content-Disposition": f"inline; filename={document.filename}"
        }
    )

@router.get("/{document_id}")
async def get_document(document_id: str, db: Session = Depends(get_db)):
    """Get document by ID."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": document.id,
        "filename": document.filename,
        "filetype": document.filetype,
        "filesize": document.filesize,
        "uploaded_at": document.uploaded_at.isoformat() if document.uploaded_at else None,
        "google_drive_id": document.google_drive_id
    }

@router.get("/{document_id}/download")
async def download_document(document_id: str, format: str = "docx", db: Session = Depends(get_db)):
    """Download document in specified format."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if document.google_drive_id and format in ['docx', 'pdf']:
        # Download from Google Drive
        file_data = get_google_drive_service().download_document(document.google_drive_id, format)
        if file_data:
            mime_types = {
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'pdf': 'application/pdf'
            }
            return Response(
                content=file_data,
                media_type=mime_types[format],
                headers={"Content-Disposition": f"attachment; filename={document.filename.split('.')[0]}.{format}"}
            )
    
    # Fallback to original file
    return Response(
        content=document.file_data,
        media_type=document.filetype,
        headers={"Content-Disposition": f"attachment; filename={document.filename}"}
    )

@router.get("/{document_id}/text")
async def get_document_text(document_id: str, db: Session = Depends(get_db)):
    """Get extracted text content from document for AI analysis."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Extract text content from the document
    text_content = extract_text_from_document(
        document.file_data, 
        document.filetype, 
        document.filename
    )
    
    return {
        "document_id": document.id,
        "filename": document.filename,
        "filetype": document.filetype,
        "text_content": text_content,
        "word_count": len(text_content.split()) if text_content else 0,
        "char_count": len(text_content) if text_content else 0
    }

@router.delete("/{document_id}")
async def delete_document(document_id: str, db: Session = Depends(get_db)):
    """Delete a single document from both database and Google Drive."""
    try:
        # Get document from database
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        google_drive_id = document.google_drive_id
        filename = document.filename
        
        # Delete from Google Drive if it exists there
        google_drive_success = True
        if google_drive_id:
            google_drive_service = get_google_drive_service()
            if google_drive_service:
                google_drive_success = google_drive_service.delete_document(google_drive_id)
                if not google_drive_success:
                    print(f"⚠️ Warning: Failed to delete {filename} from Google Drive")
            else:
                print("⚠️ Warning: Google Drive service not available")
        
        # Delete from database
        db.delete(document)
        db.commit()
        
        return {
            "message": f"Document '{filename}' deleted successfully",
            "document_id": document_id,
            "deleted_from_database": True,
            "deleted_from_google_drive": google_drive_success if google_drive_id else None,
            "google_drive_id": google_drive_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"❌ Delete failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Delete failed: {str(e)}")

@router.delete("")
async def delete_all_documents(confirm: bool = False, db: Session = Depends(get_db)):
    """Delete all documents from both database and Google Drive."""
    if not confirm:
        raise HTTPException(
            status_code=400, 
            detail="Deletion not confirmed. Add ?confirm=true to delete all documents."
        )
    
    try:
        # Get all documents
        documents = db.query(Document).all()
        
        if not documents:
            return {
                "message": "No documents found to delete",
                "deleted_count": 0
            }
        
        # Collect Google Drive IDs
        google_drive_ids = [doc.google_drive_id for doc in documents if doc.google_drive_id]
        document_names = [doc.filename for doc in documents]
        
        # Delete from Google Drive
        google_drive_results = {}
        if google_drive_ids:
            google_drive_service = get_google_drive_service()
            if google_drive_service:
                google_drive_results = google_drive_service.delete_multiple_documents(google_drive_ids)
            else:
                print("⚠️ Warning: Google Drive service not available")
        
        # Delete all from database
        deleted_count = len(documents)
        for document in documents:
            db.delete(document)
        
        db.commit()
        
        # Count successful Google Drive deletions
        successful_google_drive_deletes = sum(1 for success in google_drive_results.values() if success)
        
        return {
            "message": f"Successfully deleted {deleted_count} documents",
            "deleted_count": deleted_count,
            "document_names": document_names,
            "deleted_from_database": deleted_count,
            "deleted_from_google_drive": successful_google_drive_deletes,
            "google_drive_deletion_details": google_drive_results
        }
        
    except Exception as e:
        db.rollback()
        print(f"❌ Bulk delete failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Bulk delete failed: {str(e)}")

@router.post("/delete-selected")
async def delete_selected_documents(
    document_ids: List[str], 
    db: Session = Depends(get_db)
):
    """Delete selected documents from both database and Google Drive."""
    try:
        if not document_ids:
            raise HTTPException(status_code=400, detail="No document IDs provided")
        
        # Get documents from database
        documents = db.query(Document).filter(Document.id.in_(document_ids)).all()
        
        if not documents:
            raise HTTPException(status_code=404, detail="No documents found with provided IDs")
        
        # Collect Google Drive IDs and document info
        google_drive_ids = [doc.google_drive_id for doc in documents if doc.google_drive_id]
        document_info = {doc.id: doc.filename for doc in documents}
        
        # Delete from Google Drive
        google_drive_results = {}
        if google_drive_ids:
            google_drive_service = get_google_drive_service()
            if google_drive_service:
                google_drive_results = google_drive_service.delete_multiple_documents(google_drive_ids)
            else:
                print("⚠️ Warning: Google Drive service not available")
        
        # Delete from database
        deleted_count = 0
        for document in documents:
            db.delete(document)
            deleted_count += 1
        
        db.commit()
        
        # Count successful Google Drive deletions
        successful_google_drive_deletes = sum(1 for success in google_drive_results.values() if success)
        
        return {
            "message": f"Successfully deleted {deleted_count} selected documents",
            "deleted_count": deleted_count,
            "document_info": document_info,
            "deleted_from_database": deleted_count,
            "deleted_from_google_drive": successful_google_drive_deletes,
            "google_drive_deletion_details": google_drive_results
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"❌ Selected delete failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Selected delete failed: {str(e)}")
